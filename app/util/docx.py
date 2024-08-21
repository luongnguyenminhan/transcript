from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE

class DocxExporter:
    def __init__(self, filename):
        self.doc = Document()
        self.styles = self.doc.styles
        self.filename = filename
        self.init_styles()

    def init_styles(self):
        self._add_normal_style()
        self._add_bullet_style()
        self._add_heading_styles()

    def _add_normal_style(self):
        style = self._get_or_add_style('Normal', WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(11)

    def _add_bullet_style(self):
        style = self._get_or_add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.25)
        style.paragraph_format.first_line_indent = Inches(-0.25)

    def _add_heading_styles(self):
        for i in range(1, 5):
            style = self._get_or_add_style(f'Heading {i}', WD_STYLE_TYPE.PARAGRAPH)
            style.font.size = Pt(16 - 2 * i)
            style.font.color.rgb = RGBColor(68, 114, 196)
            style.font.bold = True

    def _get_or_add_style(self, name, style_type):
        return self.styles[name] if name in self.styles else self.styles.add_style(name, style_type)

    def _create_table(self, table_rows):
        table = self.doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        table.style = 'Table Grid'
        for i, row in enumerate(table_rows):
            for j, cell in enumerate(row):
                table.cell(i, j).text = cell
                if i == 0:
                    for paragraph in table.cell(i, j).paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

    def _add_paragraph(self, line, style, bold_parts):
        p = self.doc.add_paragraph(style=style)
        for i, part in enumerate(bold_parts):
            run = p.add_run(part)
            if i % 2 == 1:
                run.bold = True

    def export(self, prd):
        for _, content in prd.items():
            lines = content.split('\n')
            table_rows = []

            for line in lines:
                line = line.strip()
                if line.startswith('|') and line.endswith('|'):
                    table_rows.append([cell.strip() for cell in line.strip('|').split('|')])
                elif table_rows:
                    self._create_table(table_rows)
                    table_rows = []
                else:
                    self._handle_line(line)

            if table_rows:
                self._create_table(table_rows)

        self.doc.save(self.filename)

    def _handle_line(self, line):
        if line.startswith('### '):
            self.doc.add_heading(line.lstrip('#').strip(), level=2)
        elif line.startswith('#### '):
            self.doc.add_heading(line.lstrip('#').strip(), level=3)
        elif line.startswith('* '):
            self._add_paragraph(line, 'List Bullet', line.split('**'))
        else:
            self._add_paragraph(line, 'Normal', line.split('**'))
